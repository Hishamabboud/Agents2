#!/usr/bin/env python3
"""
tailor.py - Resume & Cover Letter Generator

For each qualified job in scored-jobs.json, generates a tailored resume and
cover letter combined into a single .docx file:
  output/applications/{app_id}/application.docx

App ID format: {company-slug}_{role-slug}_{6-char-url-hash}
"""

import hashlib
import io
import json
import os
import re
import sys
import time
from pathlib import Path

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SCORED_JOBS_FILE = os.path.join(BASE_DIR, "data", "scored-jobs.json")
RESUME_FILE = os.path.join(BASE_DIR, "profile", "resume.md")
COVER_LETTER_TEMPLATE_FILE = os.path.join(BASE_DIR, "profile", "cover-letter-template.md")
APPLICATIONS_OUTPUT_DIR = os.path.join(BASE_DIR, "output", "applications")

STATE_HINTS = {
    "montana": "MT", "mt": "MT", "missoula": "MT", "bozeman": "MT", "billings": "MT",
    "california": "CA", "ca": "CA", "stanford": "CA", "los angeles": "CA",
    "san francisco": "CA", "san diego": "CA", "pasadena": "CA",
    "missouri": "MO", "mo": "MO", "columbia": "MO", "st. louis": "MO", "kansas city": "MO",
    "kansas": "KS", "ks": "KS",
    "tennessee": "TN", "tn": "TN", "nashville": "TN",
    "colorado": "CO", "co": "CO", "denver": "CO", "boulder": "CO", "aurora": "CO",
    "maryland": "MD", "md": "MD", "baltimore": "MD", "bethesda": "MD",
    "michigan": "MI", "mi": "MI", "ann arbor": "MI",
    "minnesota": "MN", "mn": "MN", "rochester": "MN", "minneapolis": "MN",
    "new york": "NY", "ny": "NY",
    "washington": "WA", "wa": "WA", "seattle": "WA",
    "massachusetts": "MA", "ma": "MA", "boston": "MA", "cambridge": "MA",
    "illinois": "IL", "il": "IL", "chicago": "IL",
    "texas": "TX", "tx": "TX", "austin": "TX", "houston": "TX", "dallas": "TX",
    "remote": "REMOTE",
}


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------

def load_json(path: str, default=None):
    if default is None:
        default = []
    if not os.path.exists(path):
        return default
    try:
        with open(path) as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError):
        return default


def read_file(path: str) -> str:
    try:
        with open(path) as f:
            return f.read()
    except IOError:
        return ""


def slugify(text: str, max_len: int = 30) -> str:
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    return text[:max_len]


def detect_state(location: str, url: str) -> str:
    combined = (location + " " + url).lower()
    for hint, code in STATE_HINTS.items():
        if hint in combined:
            return code
    return "USA"


def make_app_id(job: dict) -> str:
    company_slug = slugify(job.get("company", "unknown"), max_len=20)
    role_slug = slugify(job.get("title", "role"), max_len=30)
    url_hash = hashlib.md5(job.get("url", "").encode()).hexdigest()[:6]
    return f"{company_slug}_{role_slug}_{url_hash}"


def clean_dashes(text: str) -> str:
    """Replace em/en dashes used inline with commas. Leave markdown --- separators alone."""
    # Em dash and en dash within a line (not at line boundaries)
    text = re.sub(r"(?<=[^\n])\s*\u2014\s*(?=[^\n])", ", ", text)   # —
    text = re.sub(r"(?<=[^\n])\s*\u2013\s*(?=[^\n])", ", ", text)   # –
    # Inline double-dash used as em dash (word--word) but NOT standalone --- separator lines
    text = re.sub(r"(?<=[a-zA-Z0-9])\s*--\s*(?=[a-zA-Z])", ", ", text)
    # Clean up double commas or comma after sentence-ending punctuation
    text = re.sub(r",\s*,", ",", text)
    text = re.sub(r"\.\s*,", ".", text)
    return text


# ---------------------------------------------------------------------------
# Claude API
# ---------------------------------------------------------------------------

def _get_anthropic_client():
    import anthropic
    if os.environ.get("ANTHROPIC_API_KEY"):
        return anthropic.Anthropic()
    token_file = os.environ.get("CLAUDE_SESSION_INGRESS_TOKEN_FILE", "")
    if token_file and os.path.exists(token_file):
        try:
            with open(token_file) as f:
                token = f.read().strip()
            if token:
                return anthropic.Anthropic(auth_token=token)
        except IOError:
            pass
    return anthropic.Anthropic()


def call_claude(prompt: str) -> str:
    try:
        client = _get_anthropic_client()
        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
        )
        return message.content[0].text
    except Exception as e:
        print(f"  [ERROR] Claude unavailable: {e}", file=sys.stderr)
        return ""


# ---------------------------------------------------------------------------
# Claude prompts
# ---------------------------------------------------------------------------

def generate_tailored_resume(job: dict, resume: str) -> str:
    job_desc = job.get("description", "")[:3000]
    title = job.get("title", "")
    company = job.get("company", "")

    prompt = f"""You are a professional resume writer. Tailor the resume below for this specific job. Keep it concise and to one page.

JOB TITLE: {title}
COMPANY: {company}
JOB DESCRIPTION (excerpt):
{job_desc}

ORIGINAL RESUME:
{resume}

INSTRUCTIONS:
1. Mirror EXACT keywords from the job description for ATS matching.
2. Reorganize Technical Skills into 4 to 6 labeled categories relevant to this role; put the most relevant first.
3. Rewrite bullet points to emphasize relevant experience using job description language.
4. STRICT length limits: maximum 10 bullets for Tekkii, maximum 6 bullets for Allied Engineering. Choose the most relevant ones only.
5. Each bullet: one tight sentence, maximum two lines. No long explanatory suffixes.
6. Do NOT fabricate any experience, skills, or qualifications not in the original.
7. Do NOT mention H-1B, visa status, or sponsorship.
8. Do NOT use em dashes (the long dash: \u2014) or en dashes (\u2013). Use commas, semicolons, colons, or full stops instead.
9. Do NOT use markdown horizontal rules (---). Use only headings (##, ###) to separate sections.
10. Keep Mohamad's real contact info and education unchanged.
11. Output ONLY the tailored resume in plain markdown, no commentary.

OUTPUT: Compact tailored resume in markdown."""
    result = call_claude(prompt)
    return clean_dashes(result)


def generate_cover_letter(job: dict, resume: str, template: str) -> str:
    job_desc = job.get("description", "")[:3000]
    title = job.get("title", "")
    company = job.get("company", "")
    url = job.get("url", "")

    is_healthcare = any(
        kw in (job_desc + company + title).lower()
        for kw in ["hospital", "health", "medical", "clinic", "patient", "clinical"]
    )

    opening_guidance = (
        "Your brother is a physician and your sister is a psychologist. "
        "After the opening sentence about the role, naturally weave in that growing up around healthcare conversations "
        "shaped how you think about technology work and why accurate data matters."
    ) if is_healthcare else (
        "After the opening sentence about the role, connect authentically to what this organisation does "
        "and why their mission personally resonates with you."
    )

    prompt = f"""You are writing a cover letter for Mohamad Abboud. Write in simple, warm, human language. No corporate jargon.

JOB TITLE: {title}
COMPANY: {company}
JOB URL: {url}
JOB DESCRIPTION (excerpt):
{job_desc}

CANDIDATE RESUME:
{resume}

COVER LETTER TEMPLATE:
{template}

INSTRUCTIONS:
1. Write exactly 4 paragraphs plus a closing line. Total length 480 to 580 words.
2. PARAGRAPH 1 (opening): The very first sentence must be a genuine, specific statement about why Mohamad is applying for THIS role at THIS organisation and what about their mission or work resonates with him. Use a structure like: "I am applying for the [role] at [company] because [specific reason about their work or mission]." {opening_guidance} Do NOT open with a statement about family or personal background before establishing why you want this specific job. 3 to 4 sentences total.
3. PARAGRAPH 2 (Tekkii): Walk through specific relevant work at Tekkii in detail. Mirror job description language. Mention concrete projects where relevant: the Kansas City Drayage EDI system, KidsCloset platform transformation, KC Sports Directory. 4 to 5 sentences.
4. PARAGRAPH 3 (Allied Engineering and education): Describe Allied Engineering Group work, the SQL Server and financial data experience across 400+ institutions in 40+ countries. If relevant to this role, mention the master's thesis pneumonia detection app built in Python. 4 to 5 sentences.
5. PARAGRAPH 4 (closing): Why specifically drawn to this organisation and role. What excites you about contributing there. Do NOT mention current city, current state, or willingness to relocate. 3 to 4 sentences.
6. Closing line: "Thank you for considering my application. I would love to discuss how my [relevant skills] can contribute to [company]'s [mission or goals]."
7. Sign off: Sincerely, Mohamad Abboud
8. Do NOT use em dashes (the long dash: \u2014) or en dashes (\u2013). Use commas, semicolons, colons, or full stops instead.
9. Do NOT mention H-1B, visa status, OPT, or sponsorship.
10. Output ONLY the cover letter text, no commentary.

OUTPUT: The complete cover letter."""
    result = call_claude(prompt)
    return clean_dashes(result)


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------

def _add_rich_paragraph(doc, text: str, style: str = None):
    """Add a paragraph handling inline **bold** markers."""
    from docx.shared import Pt
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part:
            p.add_run(part)
    return p


def build_docx(title: str, company: str, url: str,
               cover_letter: str, resume_md: str) -> bytes:
    """Build a single .docx: job ref + cover letter (page 1) + resume (page 2)."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Job reference header ──────────────────────────────────────────────
    ref = doc.add_paragraph()
    ref.paragraph_format.space_after = Pt(2)
    r = ref.add_run(f"{title} at {company}")
    r.bold = True
    r.font.size = Pt(11)

    url_p = doc.add_paragraph()
    url_p.paragraph_format.space_after = Pt(14)
    ur = url_p.add_run(url)
    ur.font.size = Pt(10)
    ur.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    # Thin horizontal rule (simulate with underscores)
    rule = doc.add_paragraph("_" * 80)
    rule.paragraph_format.space_after = Pt(12)
    rule.runs[0].font.size = Pt(6)

    # ── Cover letter ─────────────────────────────────────────────────────
    cl_heading = doc.add_paragraph()
    cl_heading.paragraph_format.space_after = Pt(10)
    clr = cl_heading.add_run("COVER LETTER")
    clr.bold = True
    clr.font.size = Pt(11)

    for para in cover_letter.strip().split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(0)

    # ── Page break before resume ──────────────────────────────────────────
    doc.add_page_break()

    # ── Resume ───────────────────────────────────────────────────────────
    lines = resume_md.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            h.paragraph_format.space_after = Pt(4)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            h.paragraph_format.space_after = Pt(4)
        elif line.startswith("### "):
            # Job title lines — bold paragraph, no heading style
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line[4:])
            r.bold = True
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            # Handle inline bold within bullets
            parts = re.split(r"(\*\*[^*]+\*\*)", line[2:])
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                elif part:
                    p.add_run(part)
        elif line.startswith("**") or "**" in line:
            _add_rich_paragraph(doc, line)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("=== Phase 3: Resume & Cover Letter Generation ===", flush=True)
    print("Output: output/applications/{app_id}/application.docx", flush=True)

    scored_jobs = load_json(SCORED_JOBS_FILE)
    qualified = [j for j in scored_jobs if j.get("status") == "qualified"]

    if not qualified:
        print("[WARN] No qualified jobs found. Run score.py first.")
        return 0

    resume = read_file(RESUME_FILE)
    template = read_file(COVER_LETTER_TEMPLATE_FILE)

    if not resume:
        print("[ERROR] Could not read profile/resume.md")
        return 1

    Path(APPLICATIONS_OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    generated = 0

    for i, job in enumerate(scored_jobs):
        if job.get("status") != "qualified":
            continue

        title = job.get("title", "unknown")
        company = job.get("company", "unknown")

        app_id = job.get("app_id") or make_app_id(job)
        app_dir = os.path.join(APPLICATIONS_OUTPUT_DIR, app_id)
        docx_path = os.path.join(app_dir, "application.docx")

        # Skip if already generated
        if os.path.exists(docx_path):
            print(f"  SKIP (already generated): [{app_id}] {title} @ {company}")
            scored_jobs[i]["app_id"] = app_id
            scored_jobs[i]["app_dir"] = app_dir
            scored_jobs[i]["docx_file"] = docx_path
            continue

        print(f"\n  [{app_id}]")
        print(f"  {title} @ {company} (score: {job.get('score')})")
        Path(app_dir).mkdir(parents=True, exist_ok=True)

        # Generate tailored resume
        print("    -> Tailoring resume ...", flush=True)
        tailored_resume = generate_tailored_resume(job, resume)
        if not tailored_resume:
            print("    [ERROR] Failed to generate resume")
            continue

        # Generate cover letter
        print("    -> Writing cover letter ...", flush=True)
        cover_letter = generate_cover_letter(job, resume, template)
        if not cover_letter:
            print("    [ERROR] Failed to generate cover letter")
            continue

        # Build combined .docx
        print("    -> Building application.docx ...", flush=True)
        try:
            docx_bytes = build_docx(
                title=title,
                company=company,
                url=job.get("url", ""),
                cover_letter=cover_letter,
                resume_md=tailored_resume,
            )
            with open(docx_path, "wb") as f:
                f.write(docx_bytes)
            print(f"    -> application.docx saved ({len(docx_bytes)//1024}KB)")
        except Exception as e:
            print(f"    [ERROR] Failed to build docx: {e}", file=sys.stderr)
            continue

        scored_jobs[i]["app_id"] = app_id
        scored_jobs[i]["app_dir"] = app_dir
        scored_jobs[i]["docx_file"] = docx_path
        generated += 1

        if i < len(scored_jobs) - 1:
            time.sleep(3)

    # Save updated records
    with open(SCORED_JOBS_FILE, "w") as f:
        json.dump(scored_jobs, f, indent=2)

    print(f"\n[DONE] Generated materials for {generated} job(s).")
    print("       Application packages in: output/applications/")
    return 0


if __name__ == "__main__":
    sys.exit(main())
