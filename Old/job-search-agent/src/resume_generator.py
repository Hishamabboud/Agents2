import json
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import config
import logging

logger = logging.getLogger(__name__)


class ResumeGenerator:
    def __init__(self):
        self.client = Anthropic(api_key=config.ANTHROPIC_API_KEY)

    async def generate_ats_resume(self, job_evaluation, master_resume):
        """Generate ATS-optimized resume for specific job"""
        prompt = f"""Create an ATS-optimized one-page resume for this job posting.

RULES:
- Mirror job posting keywords (especially in skills section)
- Reorder skills to match their requirements first
- Keep all content truthful (only use real experience from master resume)
- Maintain one-page format (strict)
- Quantify achievements where possible
- Use action verbs from job posting
- DO NOT add experience candidate doesn't have
- DO NOT exaggerate or lie

JOB POSTING:
Title: {job_evaluation['job_title']}
Company: {job_evaluation['company']}
Key Requirements: {job_evaluation.get('tech_stack_requirements', 'See description')}

MASTER RESUME:
{master_resume}

Return resume in markdown format with these sections:
- Header (Name, Contact)
- Technical Skills (reordered to match job)
- Experience (bullet points tailored to emphasize relevant work)
- Education

Focus on making it ATS-friendly while highlighting fit for this specific role."""

        try:
            message = self.client.messages.create(
                model=config.CLAUDE_MODEL,
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}]
            )

            resume_markdown = message.content[0].text

            # Remove markdown code blocks if present
            if "```markdown" in resume_markdown:
                resume_markdown = resume_markdown.split("```markdown")[1].split("```")[0]
            elif "```" in resume_markdown:
                resume_markdown = resume_markdown.split("```")[1].split("```")[0]

            return resume_markdown.strip()

        except Exception as e:
            logger.error(f"Error generating resume: {e}")
            return None

    def convert_to_docx(self, markdown_resume, company_name):
        """Convert markdown resume to professional .docx"""
        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Parse markdown and add to document
            lines = markdown_resume.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Header (name)
                if line.startswith('# '):
                    p = doc.add_paragraph(line[2:])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.runs[0]
                    run.font.size = Pt(16)
                    run.bold = True

                # Contact info
                elif '|' in line and '@' in line:
                    p = doc.add_paragraph(line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.runs[0]
                    run.font.size = Pt(10)

                # Section headers
                elif line.startswith('## '):
                    p = doc.add_paragraph(line[3:])
                    run = p.runs[0]
                    run.font.size = Pt(12)
                    run.bold = True

                # Bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')

                # Regular paragraphs
                else:
                    doc.add_paragraph(line)

            # Save
            safe_name = company_name.replace(' ', '_').replace('/', '_').replace('\\', '_')
            filename = f"{config.APPLICATIONS_DIR}/{safe_name}_resume.docx"
            doc.save(filename)
            logger.info(f"Saved resume: {filename}")

            return filename

        except Exception as e:
            logger.error(f"Error converting to docx: {e}")
            return None
